"""
Parse the main historical catalog: Seismic Data of Bangladesh-2023x.docx
Coverage: 1918–2022 (11 tables, each ~60-240 rows)

Schema (raw):
  Col 0: Date             (DD-MM-YYYY)
  Col 1: Time Hrs (BST)
  Col 2: Time Mts (BST)
  Col 3: Time Sec (BST)
  Col 4: Lat Degrees
  Col 5: Lat Minutes
  Col 6: Lon Degrees
  Col 7: Lon Minutes
  Col 8: Magnitude (Richter)
  Col 9: Intensity label

Outputs:
  data_intermediate/parsed_main_catalog.csv
"""

import re
import sys
from pathlib import Path

import numpy as np
import pandas as pd
import docx

PROJECT_ROOT = Path(__file__).resolve().parents[2]
DATA_RAW      = PROJECT_ROOT / "data_raw"
OUT_DIR       = PROJECT_ROOT / "data_intermediate"
OUT_DIR.mkdir(parents=True, exist_ok=True)

SOURCE_FILE = DATA_RAW / "Seismic Data of Bangladesh-2023x.docx"

# Number of header rows to skip in each table (always 3 rows in this file)
N_HEADER_ROWS = 3


def cell_text(cell) -> str:
    return " ".join(cell.text.split())


def parse_date(raw: str) -> tuple[str, list[str]]:
    """
    Parse date string. Main catalog uses DD-MM-YYYY (hyphens) for most tables,
    but Table 10 uses DD/MM/YYYY (slashes). Both are handled here.
    Returns ISO date string and a list of flag strings.
    """
    flags = []
    raw = raw.strip()

    # Attempt DD-MM-YYYY or DD/MM/YYYY (allow either separator)
    m = re.fullmatch(r"(\d{1,2})[-/](\d{1,2})[-/](\d{4})", raw)
    if m:
        d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if mo > 12:
            # Possibly MM-DD-YYYY? Flag it.
            flags.append(f"month_ambiguous:{raw}")
        try:
            dt = pd.Timestamp(year=y, month=mo, day=d)
            return dt.strftime("%Y-%m-%d"), flags
        except Exception:
            flags.append(f"date_parse_error:{raw}")
            return pd.NA, flags

    # Attempt to fix OCR/typo dates: e.g. "26-012014" → "26-01-2014"
    m2 = re.fullmatch(r"(\d{1,2})[-/](\d{2})(\d{4})", raw)
    if m2:
        d, mo, y = int(m2.group(1)), int(m2.group(2)), int(m2.group(3))
        flags.append(f"date_typo_corrected:{raw}")
        try:
            dt = pd.Timestamp(year=y, month=mo, day=d)
            return dt.strftime("%Y-%m-%d"), flags
        except Exception:
            flags.append(f"date_parse_error:{raw}")
            return pd.NA, flags

    flags.append(f"date_format_unknown:{raw}")
    return pd.NA, flags


def parse_time_bst(hrs_raw: str, mts_raw: str, sec_raw: str) -> tuple[str, bool]:
    """
    Parse BST time fields. Returns 'HH:MM:SS' string or NA, and a missing flag.
    '-' indicates missing time in the early historical record.
    """
    if all(v.strip() in ("-", "", "—") for v in [hrs_raw, mts_raw, sec_raw]):
        return pd.NA, True  # time missing (common pre-1984)

    def to_int(s: str) -> int | None:
        s = s.strip().replace("-", "").replace("—", "")
        if not s:
            return None
        try:
            return int(float(s))
        except ValueError:
            return None

    h = to_int(hrs_raw)
    m = to_int(mts_raw)
    s = to_int(sec_raw)

    if h is None or m is None or s is None:
        return pd.NA, True

    # Basic sanity
    if not (0 <= h <= 23 and 0 <= m <= 59 and 0 <= s <= 59):
        return pd.NA, True

    return f"{h:02d}:{m:02d}:{s:02d}", False


def parse_coord(deg_raw: str, mts_raw: str, hemisphere: str = "") -> tuple[float | None, list[str]]:
    """
    Convert degrees + decimal minutes to decimal degrees.
    hemisphere: 'N'/'S' for lat, 'E'/'W' for lon (sign convention).
    """
    flags = []

    def to_float(s: str) -> float | None:
        s = s.strip().replace(",", ".")
        if not s or s in ("-", "—", ""):
            return None
        try:
            return float(s)
        except ValueError:
            flags.append(f"coord_parse_error:{s}")
            return None

    deg = to_float(deg_raw)
    mts = to_float(mts_raw)

    if deg is None or mts is None:
        flags.append("coord_missing")
        return None, flags

    decimal = deg + mts / 60.0

    # For southern / western, negate
    if hemisphere in ("S", "W"):
        decimal = -decimal

    return decimal, flags


def parse_magnitude(raw: str) -> tuple[float | None, list[str]]:
    """Parse magnitude string, return float or None with flags."""
    flags = []
    raw = raw.strip().replace(",", ".")
    if not raw or raw in ("-", "—"):
        flags.append("magnitude_missing")
        return None, flags
    try:
        val = float(raw)
        if val < 0 or val > 10:
            flags.append(f"magnitude_suspect:{val}")
        return val, flags
    except ValueError:
        flags.append(f"magnitude_parse_error:{raw}")
        return None, flags


def extract_table(table, table_idx: int) -> pd.DataFrame:
    """
    Extract a single docx table into a DataFrame.
    The first 3 rows are headers and are skipped.
    """
    rows = []
    all_flags = []

    for row_idx, row in enumerate(table.rows):
        if row_idx < N_HEADER_ROWS:
            continue  # skip header rows

        cells = [cell_text(c) for c in row.cells]
        if len(cells) < 10:
            all_flags.append(f"table{table_idx}_row{row_idx}: too few cells ({len(cells)})")
            continue

        # Unpack raw cells
        date_raw  = cells[0]
        hrs_raw   = cells[1]
        mts_raw   = cells[2]
        sec_raw   = cells[3]
        lat_d_raw = cells[4]
        lat_m_raw = cells[5]
        lon_d_raw = cells[6]
        lon_m_raw = cells[7]
        mag_raw   = cells[8]
        intensity = cells[9] if len(cells) > 9 else ""

        row_flags = []

        # Parse date
        date_iso, df = parse_date(date_raw)
        row_flags.extend(df)

        # Parse time (BST)
        time_bst, time_missing = parse_time_bst(hrs_raw, mts_raw, sec_raw)
        if time_missing:
            row_flags.append("time_missing_historical")

        # Parse coordinates (all lat/lon are N/E in this catalog)
        lat, lf = parse_coord(lat_d_raw, lat_m_raw, hemisphere="N")
        row_flags.extend(lf)
        lon, lnf = parse_coord(lon_d_raw, lon_m_raw, hemisphere="E")
        row_flags.extend(lnf)

        # Parse magnitude
        mag, mf = parse_magnitude(mag_raw)
        row_flags.extend(mf)

        rows.append({
            "source_table"    : table_idx,
            "date_iso"        : date_iso,
            "time_bst"        : time_bst,
            "time_missing"    : time_missing,
            "lat_deg_raw"     : lat_d_raw,
            "lat_mts_raw"     : lat_m_raw,
            "lon_deg_raw"     : lon_d_raw,
            "lon_mts_raw"     : lon_m_raw,
            "latitude"        : lat,
            "longitude"       : lon,
            "magnitude"       : mag,
            "magnitude_raw"   : mag_raw,
            "intensity_raw"   : intensity,
            "date_raw"        : date_raw,
            "parse_flags"     : "|".join(row_flags) if row_flags else "",
        })

    return pd.DataFrame(rows), all_flags


def main():
    print(f"Parsing: {SOURCE_FILE.name}")
    doc = docx.Document(str(SOURCE_FILE))
    print(f"  Tables found: {len(doc.tables)}")

    all_dfs = []
    all_warnings = []

    for i, table in enumerate(doc.tables):
        df, warnings = extract_table(table, i)
        if warnings:
            all_warnings.extend(warnings)
        print(f"  Table {i:>2}: {len(df)} data rows extracted")
        all_dfs.append(df)

    # Concatenate all tables
    df_all = pd.concat(all_dfs, ignore_index=True)

    # Add metadata
    df_all["source_file"]   = "Seismic Data of Bangladesh-2023x.doc"
    df_all["source_period"] = "1918-2022"
    df_all["time_zone_raw"] = "BST"  # Bangladesh Standard Time = UTC+6

    # Derive year/month/day
    df_all["year"]  = pd.to_datetime(df_all["date_iso"], errors="coerce").dt.year
    df_all["month"] = pd.to_datetime(df_all["date_iso"], errors="coerce").dt.month
    df_all["day"]   = pd.to_datetime(df_all["date_iso"], errors="coerce").dt.day
    df_all["decade"] = (df_all["year"] // 10 * 10).astype("Int64")

    # Map intensity labels to numeric (for reference only — not a standardised scale)
    intensity_map = {
        "major":      8, "Major":      8,
        "strong":     6, "Strong":     6,
        "moderate":   5, "Moderate":   5,
        "light":      4, "Light":      4,
        "minor":      3, "Minor":      3,
        "very minor": 2, "Very Minor": 2,
    }
    df_all["intensity_label"] = df_all["intensity_raw"].str.strip()
    df_all["intensity_numeric"] = df_all["intensity_label"].map(
        lambda x: next((v for k, v in intensity_map.items() if k.lower() == x.lower()), None)
    )

    # Sanity checks
    n_total      = len(df_all)
    n_no_date    = df_all["date_iso"].isna().sum()
    n_no_coord   = (df_all["latitude"].isna() | df_all["longitude"].isna()).sum()
    n_no_mag     = df_all["magnitude"].isna().sum()
    n_no_time    = df_all["time_missing"].sum()

    print(f"\n  Total rows       : {n_total}")
    print(f"  Missing date     : {n_no_date}")
    print(f"  Missing coords   : {n_no_coord}")
    print(f"  Missing magnitude: {n_no_mag}")
    print(f"  Missing time     : {n_no_time} ({n_no_time/n_total*100:.1f}%)")
    print(f"  Year range       : {df_all['year'].min()} – {df_all['year'].max()}")
    print(f"  Magnitude range  : {df_all['magnitude'].min()} – {df_all['magnitude'].max()}")

    if all_warnings:
        print(f"\n  WARNINGS ({len(all_warnings)}):")
        for w in all_warnings[:20]:
            print(f"    {w}")

    # Save
    out_path = OUT_DIR / "parsed_main_catalog.csv"
    df_all.to_csv(out_path, index=False, encoding="utf-8")
    print(f"\n  Saved → {out_path.relative_to(PROJECT_ROOT)}")
    print(f"  Columns: {list(df_all.columns)}")

    return df_all


if __name__ == "__main__":
    main()
