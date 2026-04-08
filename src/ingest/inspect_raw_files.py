"""
Step 1: Raw file inspection script.
Reads all source DOCX files, extracts tables and text, and saves
a plain-text dump of each document for manual review.

Outputs:
  data_intermediate/inspection/<stem>_tables.txt   -- all tables as CSV blocks
  data_intermediate/inspection/<stem>_text.txt     -- all paragraph text
  data_intermediate/inspection/file_inventory.json -- machine-readable summary
"""

import json
import os
import sys
from pathlib import Path

import docx

# ── Paths ──────────────────────────────────────────────────────────────────────
PROJECT_ROOT = Path(__file__).resolve().parents[2]
DATA_RAW     = PROJECT_ROOT / "data_raw"
OUT_DIR      = PROJECT_ROOT / "data_intermediate" / "inspection"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# All source docx files (the .doc was already converted to .docx)
SOURCE_FILES = {
    "main_catalog": DATA_RAW / "Seismic Data of Bangladesh-2023x.docx",
    "felt_2024_2025": DATA_RAW / "Bangladesh felt Data_January 2024-24 January 2025.docx",
    "felt_2025": DATA_RAW / "Bangladesh fell Data 2025(January-(August).docx",
    "monthly_2023_2024": DATA_RAW / "মাসিক ডাটা ২৩-২৪.docx",
}


def cell_text(cell) -> str:
    """Return cleaned text from a single table cell."""
    return " ".join(cell.text.split())


def table_to_rows(table) -> list[list[str]]:
    """Convert a docx table to a list of row-lists of strings."""
    rows = []
    for row in table.rows:
        rows.append([cell_text(c) for c in row.cells])
    return rows


def rows_to_csv_block(rows: list[list[str]], table_idx: int) -> str:
    """Format rows as a readable CSV-like block with a header."""
    lines = [f"--- TABLE {table_idx} ({len(rows)} rows x {len(rows[0]) if rows else 0} cols) ---"]
    for row in rows:
        lines.append(" | ".join(row))
    return "\n".join(lines)


def inspect_docx(label: str, path: Path) -> dict:
    """
    Open a docx, extract all tables and paragraph text.
    Returns a summary dict and writes text dumps to OUT_DIR.
    """
    print(f"\n{'='*70}")
    print(f"  Inspecting: {label}")
    print(f"  File:       {path.name}")
    print(f"  Size:       {path.stat().st_size / 1024:.1f} KB")
    print(f"{'='*70}")

    try:
        doc = docx.Document(str(path))
    except Exception as e:
        print(f"  ERROR opening file: {e}")
        return {"label": label, "file": str(path), "error": str(e)}

    # ── Paragraph text ─────────────────────────────────────────────────────────
    para_texts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            para_texts.append(t)

    # ── Tables ─────────────────────────────────────────────────────────────────
    table_summaries = []
    all_table_blocks = []

    for i, tbl in enumerate(doc.tables):
        rows = table_to_rows(tbl)
        n_rows = len(rows)
        n_cols = len(rows[0]) if rows else 0
        header = rows[0] if rows else []

        print(f"  Table {i:>3}: {n_rows} rows × {n_cols} cols | header: {header[:6]}")

        block = rows_to_csv_block(rows, i)
        all_table_blocks.append(block)

        table_summaries.append({
            "table_index": i,
            "n_rows": n_rows,
            "n_cols": n_cols,
            "header_sample": header,
        })

    print(f"\n  Paragraphs (non-empty): {len(para_texts)}")
    print(f"  Tables found:           {len(doc.tables)}")

    # First few paragraphs for a quick look
    if para_texts:
        print("\n  -- First 10 paragraphs --")
        for p in para_texts[:10]:
            print(f"    {p[:120]}")

    # ── Save text dumps ────────────────────────────────────────────────────────
    stem = label.replace(" ", "_")

    table_dump_path = OUT_DIR / f"{stem}_tables.txt"
    table_dump_path.write_text("\n\n".join(all_table_blocks), encoding="utf-8")

    text_dump_path = OUT_DIR / f"{stem}_text.txt"
    text_dump_path.write_text("\n".join(para_texts), encoding="utf-8")

    print(f"\n  Saved table dump → {table_dump_path.relative_to(PROJECT_ROOT)}")
    print(f"  Saved text  dump → {text_dump_path.relative_to(PROJECT_ROOT)}")

    return {
        "label": label,
        "file": path.name,
        "size_kb": round(path.stat().st_size / 1024, 1),
        "n_paragraphs": len(para_texts),
        "n_tables": len(doc.tables),
        "tables": table_summaries,
        "first_paragraphs": para_texts[:5],
    }


def main():
    inventory = {}
    for label, path in SOURCE_FILES.items():
        if not path.exists():
            print(f"\nWARNING: file not found — {path}")
            inventory[label] = {"label": label, "file": str(path), "error": "not found"}
            continue
        inventory[label] = inspect_docx(label, path)

    inv_path = OUT_DIR / "file_inventory.json"
    inv_path.write_text(json.dumps(inventory, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"\n\nInventory saved → {inv_path.relative_to(PROJECT_ROOT)}")

    # Print summary table
    print("\n" + "="*70)
    print("  FILE INVENTORY SUMMARY")
    print("="*70)
    print(f"  {'Label':<25} {'File':<45} {'Tables':>6} {'Paragraphs':>10}")
    print(f"  {'-'*25} {'-'*45} {'-'*6} {'-'*10}")
    for label, info in inventory.items():
        if "error" in info:
            print(f"  {label:<25}  ERROR: {info['error']}")
        else:
            print(f"  {label:<25} {info['file']:<45} {info['n_tables']:>6} {info['n_paragraphs']:>10}")


if __name__ == "__main__":
    main()
