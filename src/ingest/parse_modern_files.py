"""
Parse modern felt-earthquake files:
  1. Bangladesh felt Data_January 2024-24 January 2025.docx  (Jan 2024 – Jan 2025)
  2. Bangladesh fell Data 2025(January-(August).docx           (Jan 2025 – Aug 2025)
  3. মাসিক ডাটা ২৩-২৪.docx                                   (Jul 2023 – Jun 2024, monthly)

All three use UTC time and DMS-format coordinates.

The monthly file (মাসিক ডাটা) has paired tables per month:
  - Even-indexed tables (0,1,3,5,...): broader regionally detected events
  - Odd-indexed within each month: felt/nearby Bangladesh subset
  (July 2023 only has a broader table; Aug 2023 – Jun 2024 have pairs.)

Bijoy-encoded Bengali paragraph titles (decoded from ANSI patterns):
  Para 1:  Jul 2023 – Broader
  Para 2:  Aug 2023 – Broader
  Para 3:  Aug 2023 – Felt
  Para 4:  Sep 2023 – Broader
  Para 5:  Sep 2023 – Felt
  Para 6:  Oct 2023 – Broader
  Para 7:  Oct 2023 – Felt
  Para 8:  Nov 2023 – Broader
  Para 9-10: Nov 2023 – Felt (split paragraph)
  Para 11: Dec 2023 – Broader
  Para 12: Dec 2023 – Felt
  Para 13: Jan 2024 – Broader
  Para 14: Jan 2024 – Felt
  Para 15: Feb 2024 – Broader
  Para 16: Feb 2024 – Felt
  Para 17: Mar 2024 – Broader
  Para 18: Mar 2024 – Felt
  Para 19: Apr 2024 – Broader (no felt pair)
  Para 20: May 2024 – Broader
  Para 21: May 2024 – Felt
  Para 22: Jun 2024 – Broader
  Para 23: Jun 2024 – Felt

Outputs:
  data_intermediate/parsed_felt_2024_2025.csv
  data_intermediate/parsed_felt_2025.csv
  data_intermediate/parsed_monthly_2023_2024.csv
"""

import re
from pathlib import Path

import pandas as pd
import docx

PROJECT_ROOT = Path(__file__).resolve().parents[2]
DATA_RAW     = PROJECT_ROOT / "data_raw"
OUT_DIR      = PROJECT_ROOT / "data_intermediate"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# ── Coordinate parser ──────────────────────────────────────────────────────────

# Matches patterns like:
#   26° 12.30′N   /   095°29.94′ E   /   24°36.78′N   /   240 45.6′ N
# Also handles garbled variants like "240 45.6′ N" (missing degree sign)
_COORD_RE = re.compile(
    r"([+-]?\d+(?:\.\d+)?)"     # integer degrees
    r"[°o\s]+"                  # degree sign or space
    r"(\d+(?:\.\d+)?)"          # decimal minutes
    r"[′'\u2032\s]*"            # minute sign (various)
    r"([NSEW]?)",               # optional hemisphere
    re.IGNORECASE,
)


def parse_dms_coord(raw: str) -> tuple[float | None, str, list[str]]:
    """
    Parse a DMS coordinate string into decimal degrees.
    Returns (decimal_value, hemisphere_char, flags_list).
    Hemisphere is inferred from the string if present.
    """
    flags = []
    raw = raw.strip()

    if not raw or raw in ("-", "—", ""):
        flags.append("coord_empty")
        return None, "", flags

    m = _COORD_RE.search(raw)
    if not m:
        flags.append(f"coord_unrecognised:{raw}")
        return None, "", flags

    deg = float(m.group(1))
    mts = float(m.group(2))
    hemi = m.group(3).upper() if m.group(3) else ""

    # Minutes must be [0, 60)
    if mts >= 60:
        flags.append(f"minutes_out_of_range:{mts}")
        mts = mts % 60  # best-effort fix

    decimal = deg + mts / 60.0

    if hemi in ("S", "W"):
        decimal = -decimal

    return decimal, hemi, flags


def parse_latlon_pair(lat_raw: str, lon_raw: str) -> tuple[float | None, float | None, list[str]]:
    """Parse a latitude and longitude string pair."""
    flags = []

    lat, lat_hemi, lf = parse_dms_coord(lat_raw)
    flags.extend([f"lat_{f}" for f in lf])

    lon, lon_hemi, lnf = parse_dms_coord(lon_raw)
    flags.extend([f"lon_{f}" for f in lnf])

    # Cross-validate hemispheres for a Bangladesh-centric catalog
    # Latitude should be N (positive), longitude should be E (positive)
    if lat is not None and lat_hemi == "S":
        flags.append("lat_is_S_flagged")
    if lon is not None and lon_hemi == "W":
        flags.append("lon_is_W_flagged")

    return lat, lon, flags


def parse_date_slash(raw: str) -> tuple[str, list[str]]:
    """Parse DD/MM/YYYY date format."""
    flags = []
    raw = raw.strip()
    m = re.fullmatch(r"(\d{1,2})[/.](\d{1,2})[/.](\d{4})", raw)
    if m:
        d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        try:
            dt = pd.Timestamp(year=y, month=mo, day=d)
            return dt.strftime("%Y-%m-%d"), flags
        except Exception:
            flags.append(f"date_invalid:{raw}")
            return pd.NA, flags
    flags.append(f"date_format_unknown:{raw}")
    return pd.NA, flags


def parse_time_utc(raw: str) -> tuple[str, list[str]]:
    """
    Parse UTC time from strings like '21 45 51.00' or '21 45 51'.
    Returns 'HH:MM:SS' or NA.
    """
    flags = []
    raw = raw.strip()
    if not raw or raw in ("-", "—"):
        flags.append("time_missing")
        return pd.NA, flags

    # Remove trailing decimal seconds
    raw_clean = re.sub(r"\.\d+$", "", raw)
    parts = raw_clean.split()
    if len(parts) == 3:
        try:
            h, m, s = int(parts[0]), int(parts[1]), int(parts[2])
            if 0 <= h <= 23 and 0 <= m <= 59 and 0 <= s <= 59:
                return f"{h:02d}:{m:02d}:{s:02d}", flags
            else:
                flags.append(f"time_range_error:{raw}")
                return pd.NA, flags
        except ValueError:
            pass

    flags.append(f"time_format_unknown:{raw}")
    return pd.NA, flags


def parse_magnitude(raw: str) -> tuple[float | None, list[str]]:
    flags = []
    raw = raw.strip().replace(",", ".")
    if not raw or raw in ("-", "—"):
        flags.append("magnitude_missing")
        return None, flags
    try:
        val = float(raw)
        if val < 0 or val > 10:
            flags.append(f"magnitude_suspect:{val}")
        return val, flags
    except ValueError:
        flags.append(f"magnitude_parse_error:{raw}")
        return None, flags


def parse_distance(raw: str) -> tuple[float | None, list[str]]:
    flags = []
    raw = raw.strip().replace(",", "")
    if not raw or raw in ("-", "—"):
        return None, flags
    try:
        return float(raw), flags
    except ValueError:
        flags.append(f"distance_parse_error:{raw}")
        return None, flags


def cell_text(cell) -> str:
    return " ".join(cell.text.split())


# ── Parser: felt_2024_2025 and felt_2025 ──────────────────────────────────────

def parse_felt_table(table, schema_type: str = "7col") -> pd.DataFrame:
    """
    Parse a modern felt-earthquakes table.

    schema_type='7col' (2024-2025 file):
      0:Date  1:Time(UTC)  2:Lat  3:Lon  4:Magnitude  5:DistDhaka  6:Region

    schema_type='8col' (2025 file, has SL column at start):
      0:SL  1:Date  2:Time(UTC)  3:Lat  4:Lon  5:Magnitude  6:DistDhaka  7:Region
    """
    rows = []

    for row_idx, row in enumerate(table.rows):
        if row_idx < 2:  # skip 2-row merged header
            continue

        cells = [cell_text(c) for c in row.cells]

        if schema_type == "8col":
            if len(cells) < 8:
                continue
            sl       = cells[0]
            date_raw = cells[1]
            time_raw = cells[2]
            lat_raw  = cells[3]
            lon_raw  = cells[4]
            mag_raw  = cells[5]
            dist_raw = cells[6]
            region   = cells[7]
        else:  # 7col
            if len(cells) < 7:
                continue
            sl       = ""
            date_raw = cells[0]
            time_raw = cells[1]
            lat_raw  = cells[2]
            lon_raw  = cells[3]
            mag_raw  = cells[4]
            dist_raw = cells[5]
            region   = cells[6]

        row_flags = []

        date_iso, df = parse_date_slash(date_raw)
        row_flags.extend(df)

        time_utc, tf = parse_time_utc(time_raw)
        row_flags.extend(tf)

        lat, lon, cf = parse_latlon_pair(lat_raw, lon_raw)
        row_flags.extend(cf)

        mag, mf = parse_magnitude(mag_raw)
        row_flags.extend(mf)

        dist_km, disf = parse_distance(dist_raw)
        row_flags.extend(disf)

        rows.append({
            "sl"                    : sl,
            "date_iso"              : date_iso,
            "time_utc"              : time_utc,
            "lat_raw"               : lat_raw,
            "lon_raw"               : lon_raw,
            "latitude"              : lat,
            "longitude"             : lon,
            "magnitude"             : mag,
            "magnitude_raw"         : mag_raw,
            "distance_dhaka_km"     : dist_km,
            "region_raw"            : region,
            "date_raw"              : date_raw,
            "parse_flags"           : "|".join(row_flags) if row_flags else "",
        })

    return pd.DataFrame(rows)


# ── Parser: monthly_2023_2024 ─────────────────────────────────────────────────

# Table index → (month_label, catalog_type)
# Derived by decoding Bijoy-encoded Bengali paragraph titles
MONTHLY_TABLE_MAP = {
    0:  ("2023-07", "broader"),
    1:  ("2023-08", "broader"),
    2:  ("2023-08", "felt_nearby"),
    3:  ("2023-09", "broader"),
    4:  ("2023-09", "felt_nearby"),
    5:  ("2023-10", "broader"),
    6:  ("2023-10", "felt_nearby"),
    7:  ("2023-11", "broader"),
    8:  ("2023-11", "felt_nearby"),
    9:  ("2023-12", "broader"),
    10: ("2023-12", "felt_nearby"),
    11: ("2024-01", "broader"),
    12: ("2024-01", "felt_nearby"),
    13: ("2024-02", "broader"),
    14: ("2024-02", "felt_nearby"),
    15: ("2024-03", "broader"),
    16: ("2024-03", "felt_nearby"),
    17: ("2024-04", "broader"),
    18: ("2024-05", "broader"),
    19: ("2024-05", "felt_nearby"),
    20: ("2024-06", "broader"),
    21: ("2024-06", "felt_nearby"),
}


def parse_monthly_table(table, table_idx: int) -> pd.DataFrame:
    """
    Parse a monthly catalog table (8-column format with depth).
    Schema:
      0:Date  1:Time(UTC)  2:Lat  3:Lon  4:Depth  5:Magnitude  6:DistDhaka  7:Region
    """
    rows = []
    month_label, catalog_type = MONTHLY_TABLE_MAP.get(table_idx, ("unknown", "unknown"))

    for row_idx, row in enumerate(table.rows):
        if row_idx < 2:  # skip 2 header rows
            continue

        cells = [cell_text(c) for c in row.cells]
        if len(cells) < 8:
            continue

        date_raw  = cells[0]
        time_raw  = cells[1]
        lat_raw   = cells[2]
        lon_raw   = cells[3]
        depth_raw = cells[4]
        mag_raw   = cells[5]
        dist_raw  = cells[6]
        region    = cells[7]

        row_flags = []

        date_iso, df = parse_date_slash(date_raw)
        row_flags.extend(df)

        time_utc, tf = parse_time_utc(time_raw)
        row_flags.extend(tf)

        lat, lon, cf = parse_latlon_pair(lat_raw, lon_raw)
        row_flags.extend(cf)

        mag, mf = parse_magnitude(mag_raw)
        row_flags.extend(mf)

        dist_km, disf = parse_distance(dist_raw)
        row_flags.extend(disf)

        # Depth: '--' means not reported
        depth_km = None
        if depth_raw.strip() not in ("-", "--", "—", ""):
            try:
                depth_km = float(depth_raw.replace(",", "."))
            except ValueError:
                row_flags.append(f"depth_parse_error:{depth_raw}")

        rows.append({
            "date_iso"              : date_iso,
            "time_utc"              : time_utc,
            "lat_raw"               : lat_raw,
            "lon_raw"               : lon_raw,
            "latitude"              : lat,
            "longitude"             : lon,
            "depth_km"              : depth_km,
            "magnitude"             : mag,
            "magnitude_raw"         : mag_raw,
            "distance_dhaka_km"     : dist_km,
            "region_raw"            : region,
            "date_raw"              : date_raw,
            "catalog_month"         : month_label,
            "catalog_type"          : catalog_type,
            "parse_flags"           : "|".join(row_flags) if row_flags else "",
        })

    return pd.DataFrame(rows)


def add_time_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Derive year, month, day, decade from date_iso."""
    dt = pd.to_datetime(df["date_iso"], errors="coerce")
    df["year"]   = dt.dt.year
    df["month"]  = dt.dt.month
    df["day"]    = dt.dt.day
    df["decade"] = (dt.dt.year // 10 * 10).astype("Int64")
    return df


# ── Main ──────────────────────────────────────────────────────────────────────

def main():

    # 1. felt_2024_2025
    print("\n" + "="*60)
    print("Parsing: Bangladesh felt Data_January 2024-24 January 2025.docx")
    path_24_25 = DATA_RAW / "Bangladesh felt Data_January 2024-24 January 2025.docx"
    doc_24_25  = docx.Document(str(path_24_25))
    df_24_25   = parse_felt_table(doc_24_25.tables[0], schema_type="7col")
    df_24_25["source_file"]   = "Bangladesh felt Data_January 2024-24 January 2025.docx"
    df_24_25["source_period"] = "2024-01 to 2025-01"
    df_24_25["catalog_type"]  = "felt_near_bangladesh"
    df_24_25["time_zone_raw"] = "UTC"
    df_24_25 = add_time_columns(df_24_25)

    print(f"  Rows extracted: {len(df_24_25)}")
    print(f"  Year range: {df_24_25['year'].min()} – {df_24_25['year'].max()}")
    print(f"  Magnitude range: {df_24_25['magnitude'].min()} – {df_24_25['magnitude'].max()}")

    out_24_25 = OUT_DIR / "parsed_felt_2024_2025.csv"
    df_24_25.to_csv(out_24_25, index=False, encoding="utf-8")
    print(f"  Saved → {out_24_25.relative_to(PROJECT_ROOT)}")

    # 2. felt_2025
    print("\n" + "="*60)
    print("Parsing: Bangladesh fell Data 2025(January-(August).docx")
    path_2025 = DATA_RAW / "Bangladesh fell Data 2025(January-(August).docx"
    doc_2025  = docx.Document(str(path_2025))
    df_2025   = parse_felt_table(doc_2025.tables[0], schema_type="8col")
    df_2025["source_file"]   = "Bangladesh fell Data 2025(January-(August).docx"
    df_2025["source_period"] = "2025-01 to 2025-08"
    df_2025["catalog_type"]  = "felt_near_bangladesh"
    df_2025["time_zone_raw"] = "UTC"
    df_2025 = add_time_columns(df_2025)

    print(f"  Rows extracted: {len(df_2025)}")
    print(f"  Year range: {df_2025['year'].min()} – {df_2025['year'].max()}")
    print(f"  Magnitude range: {df_2025['magnitude'].min()} – {df_2025['magnitude'].max()}")

    out_2025 = OUT_DIR / "parsed_felt_2025.csv"
    df_2025.to_csv(out_2025, index=False, encoding="utf-8")
    print(f"  Saved → {out_2025.relative_to(PROJECT_ROOT)}")

    # 3. Monthly 2023-2024
    print("\n" + "="*60)
    print("Parsing: মাসিক ডাটা ২৩-২৪.docx (monthly Jul 2023 – Jun 2024)")
    path_monthly = DATA_RAW / "মাসিক ডাটা ২৩-২৪.docx"
    doc_monthly  = docx.Document(str(path_monthly))
    print(f"  Tables found: {len(doc_monthly.tables)}")

    all_monthly = []
    for i, tbl in enumerate(doc_monthly.tables):
        df_m = parse_monthly_table(tbl, i)
        month_label, ctype = MONTHLY_TABLE_MAP.get(i, ("unknown", "unknown"))
        print(f"  Table {i:>2}: {month_label} ({ctype}) → {len(df_m)} rows")
        all_monthly.append(df_m)

    df_monthly = pd.concat(all_monthly, ignore_index=True)
    df_monthly["source_file"]   = "মাসিক ডাটা ২৩-২৪.docx"
    df_monthly["source_period"] = "2023-07 to 2024-06"
    df_monthly["time_zone_raw"] = "UTC"
    df_monthly = add_time_columns(df_monthly)

    print(f"\n  Total rows: {len(df_monthly)}")
    print(f"  Broader detected rows:   {(df_monthly['catalog_type']=='broader').sum()}")
    print(f"  Felt-nearby BD rows:     {(df_monthly['catalog_type']=='felt_nearby').sum()}")
    print(f"  Year range: {df_monthly['year'].min()} – {df_monthly['year'].max()}")
    print(f"  Magnitude range: {df_monthly['magnitude'].min()} – {df_monthly['magnitude'].max()}")

    out_monthly = OUT_DIR / "parsed_monthly_2023_2024.csv"
    df_monthly.to_csv(out_monthly, index=False, encoding="utf-8")
    print(f"  Saved → {out_monthly.relative_to(PROJECT_ROOT)}")

    return df_24_25, df_2025, df_monthly


if __name__ == "__main__":
    main()
